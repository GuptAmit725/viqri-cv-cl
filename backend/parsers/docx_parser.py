"""
DOCX Parser Module
Extracts text content from Word documents
"""

from docx import Document


def parse_docx(filepath):
    """
    Parse DOCX file and extract text content
    
    Args:
        filepath (str): Path to the DOCX file
        
    Returns:
        str: Extracted text from DOCX
    """
    try:
        doc = Document(filepath)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text += " | ".join(row_text) + "\n"
            text += "\n"
        
        return text.strip()
    
    except Exception as e:
        raise Exception(f"Error parsing DOCX: {str(e)}")


def parse_docx_structured(filepath):
    """
    Parse DOCX with structure preservation
    
    Args:
        filepath (str): Path to the DOCX file
        
    Returns:
        dict: Structured document data
    """
    try:
        doc = Document(filepath)
        
        result = {
            'paragraphs': [],
            'tables': [],
            'full_text': ''
        }
        
        # Extract paragraphs with style information
        for para in doc.paragraphs:
            if para.text.strip():
                para_data = {
                    'text': para.text,
                    'style': para.style.name if para.style else 'Normal'
                }
                result['paragraphs'].append(para_data)
                result['full_text'] += para.text + "\n"
        
        # Extract tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            result['tables'].append(table_data)
        
        return result
    
    except Exception as e:
        raise Exception(f"Error parsing DOCX structure: {str(e)}")